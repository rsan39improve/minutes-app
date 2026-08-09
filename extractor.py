"""
PDF・テキスト・Wordファイル、および貼り付けテキストから本文を抽出するモジュール。
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Optional, Tuple, Union

import pdfplumber
from docx import Document

MAX_TRANSCRIPT_CHARS = 100_000


class EmptyPdfTextError(ValueError):
    """PDFから文字が抽出できなかった場合。"""


def extract_text_from_pdf(file: Union[str, Path, bytes, io.BytesIO]) -> str:
    """PDFからテキストを抽出する。文字が取れない場合は EmptyPdfTextError。"""
    if isinstance(file, (str, Path)):
        with pdfplumber.open(file) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    else:
        if isinstance(file, bytes):
            file = io.BytesIO(file)
        with pdfplumber.open(file) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]

    text = "\n".join(pages).strip()
    if not text:
        raise EmptyPdfTextError("PDFから文字を抽出できませんでした（スキャン画像の可能性）。")
    return text


def extract_text_from_txt(file: Union[str, Path, bytes]) -> str:
    """テキストファイルの内容を返す。"""
    if isinstance(file, bytes):
        for encoding in ("utf-8", "utf-8-sig", "shift-jis", "cp932"):
            try:
                return file.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file.decode("utf-8", errors="replace")

    path = Path(file)
    for encoding in ("utf-8", "utf-8-sig", "shift-jis", "cp932"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text_from_docx(file: Union[str, Path, bytes, io.BytesIO]) -> str:
    """docxから段落テキストを抽出する。"""
    if isinstance(file, (str, Path)):
        doc = Document(str(file))
    else:
        if isinstance(file, bytes):
            file = io.BytesIO(file)
        doc = Document(file)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # テーブル内テキストも拾う
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts).strip()


def extract_text_from_upload(
    file_bytes: bytes,
    filename: str,
) -> Tuple[Optional[str], Optional[str]]:
    """
    アップロードファイルからテキストを抽出する。

    Returns:
        (text, warning)
        text が None のときは抽出失敗（呼び出し側で資料なし扱い等）。
    """
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            return extract_text_from_pdf(file_bytes), None
        if name.endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
            if not text:
                return None, "Wordファイルから文字を抽出できませんでした。"
            return text, None
        # txt その他
        text = extract_text_from_txt(file_bytes).strip()
        if not text:
            return None, "テキストファイルが空です。"
        return text, None
    except EmptyPdfTextError as e:
        return None, str(e)
    except Exception as e:
        return None, f"ファイル読み込みエラー: {e}"


def normalize_transcript(text: str) -> str:
    """文字数上限チェック付きで整形する。"""
    cleaned = (text or "").strip()
    if len(cleaned) > MAX_TRANSCRIPT_CHARS:
        raise ValueError(
            f"文字起こしが長すぎます（{len(cleaned):,}文字）。"
            f"上限は {MAX_TRANSCRIPT_CHARS:,} 文字です。"
        )
    return cleaned


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("使用方法: python extractor.py <ファイルパス>")
        sys.exit(1)

    path = Path(sys.argv[1])
    data = path.read_bytes()
    text, warning = extract_text_from_upload(data, path.name)
    if warning:
        print("警告:", warning)
    print("=== 抽出テキスト ===")
    print((text or "")[:2000])
    print(f"\n... 合計 {len(text or '')} 文字")
