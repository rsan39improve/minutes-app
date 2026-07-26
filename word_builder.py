"""
python-docxで議事録Wordファイルを生成するモジュール。
「議事録ひな型.docx」の表構造に合わせた様式で出力する。
AIには触らせず、このモジュールだけが様式を制御する。
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Emu


# ===== 列幅定数（ひな型から計測）=====
COL0_WIDTH = Emu(796290)    # ラベル列（約2.1cm）
COL1_WIDTH = Emu(3947795)   # 主コンテンツ列（約10.5cm）
COL2_WIDTH = Emu(1729740)   # サブコンテンツ列（約4.6cm）

TITLE_COL0_WIDTH = Emu(1194435)   # タイトル「記録」列
TITLE_COL1_WIDTH = Emu(5286375)   # タイトル 会議名列

FONT_NAME_JP = "游明朝"


def _set_col_width(table, widths: list):
    """テーブルの各列幅を設定する。"""
    for i, col in enumerate(table.columns):
        if i < len(widths):
            col.width = widths[i]
            for cell in col.cells:
                cell.width = widths[i]


def _set_border(element, **kwargs):
    """
    セルまたはテーブルの罫線を設定する汎用関数。
    kwargs: top/bottom/left/right/insideH/insideV = (val, sz, color)
    val例: "single", "double", "nil"
    """
    tag_map = {
        "top": "w:top", "bottom": "w:bottom",
        "left": "w:left", "right": "w:right",
        "insideH": "w:insideH", "insideV": "w:insideV",
    }
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for side, spec in kwargs.items():
        tag = tag_map[side]
        child = OxmlElement(tag)
        if spec == "nil":
            child.set(qn("w:val"), "nil")
        else:
            val, sz, color = spec
            child.set(qn("w:val"), val)
            child.set(qn("w:sz"), str(sz))
            child.set(qn("w:space"), "0")
            child.set(qn("w:color"), color)
        element.append(child)


def _apply_title_table_borders(table):
    """
    ひな型「記録」帯の罫線を再現する。
    - テーブル外枠: 細い単線（single, sz=4）
    - 上下のみ二重線（double, sz=4）、左右・中央縦はなし
    """
    # テーブル全体プロパティ
    tbl = table._tbl
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tblPr = tbl.find(f"{{{ns}}}tblPr")
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    _set_border(tblBorders,
                top=("single", 4, "auto"),
                left=("single", 4, "auto"),
                bottom=("single", 4, "auto"),
                right=("single", 4, "auto"),
                insideH=("single", 4, "auto"),
                insideV=("single", 4, "auto"))
    tblPr.append(tblBorders)
    # セル余白をひな型に合わせる（左右99dxa）
    tblCellMar = OxmlElement("w:tblCellMar")
    for side in ("w:left", "w:right"):
        m = OxmlElement(side)
        m.set(qn("w:w"), "99")
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    # セル0（記録）: 上下=double、左右=nil
    tc0 = table.cell(0, 0)._tc
    tcPr0 = tc0.get_or_add_tcPr()
    b0 = OxmlElement("w:tcBorders")
    _set_border(b0,
                top=("double", 4, "auto"),
                left="nil",
                bottom=("double", 4, "auto"),
                right="nil")
    tcPr0.append(b0)

    # セル1（会議名）: 上下=double、右=nil
    tc1 = table.cell(0, 1)._tc
    tcPr1 = tc1.get_or_add_tcPr()
    b1 = OxmlElement("w:tcBorders")
    _set_border(b1,
                top=("double", 4, "auto"),
                bottom=("double", 4, "auto"),
                right="nil")
    tcPr1.append(b1)
    # 縦位置を中央揃え
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr1.append(vAlign)


def _set_font(run, size_pt: int, bold: bool = False,
              color: RGBColor = RGBColor(0, 0, 0)):
    """runのフォントを設定する。"""
    run.font.name = FONT_NAME_JP
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_NAME_JP)
    rPr.insert(0, rFonts)


def _cell_para(cell, text: str, size_pt: int = 10, bold: bool = False,
               align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """セルの最初の段落にテキストを設定する。"""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    _set_font(run, size_pt, bold=bold)


def _cell_add_para(cell, text: str, size_pt: int = 10, bold: bool = False,
                   indent_cm: float = 0.0,
                   space_before: float = 0, space_after: float = 2) -> None:
    """セルに段落を追加する。"""
    para = cell.add_paragraph()
    para.paragraph_format.left_indent = Cm(indent_cm)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    _set_font(run, size_pt, bold=bold)


def _merge_h(table, row: int, col_start: int, col_end: int):
    """行の水平方向セルをマージする。"""
    table.cell(row, col_start).merge(table.cell(row, col_end))


def _merge_full(table, row: int, n_cols: int = 3):
    """行の全列をマージする。"""
    table.cell(row, 0).merge(table.cell(row, n_cols - 1))


def build_minutes_docx(data: dict, district: str = "") -> bytes:
    """
    議事録データ（dict）からWordファイルを生成してバイト列で返す。

    Args:
        data: extract_minutes()が返す構造化データ
        district: 地区名（ページヘッダーに表示）

    Returns:
        .docxファイルのバイト列
    """
    doc = Document()

    # ===== 用紙・余白設定 =====
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.2)

    # ===== ページヘッダー（地区名）=====
    if district:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run(district)
        _set_font(header_run, 10)

    # ===== テーブル0: タイトル行 =====
    # 「記録 | ●●●打合せ」の2列形式
    title_tbl = doc.add_table(rows=1, cols=2)
    _set_col_width(title_tbl, [TITLE_COL0_WIDTH, TITLE_COL1_WIDTH])

    _apply_title_table_borders(title_tbl)

    meeting_name = data.get("meeting_name", "会議") or "会議"
    _cell_para(title_tbl.cell(0, 0), "記録", size_pt=14, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(title_tbl.cell(0, 1), meeting_name,
               size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ===== テーブル1: 情報テーブル =====
    # 行構成: 開催日時 / 場所 / 出席者×n / 資料 / 議事次第 / 議事内容ヘッダー / 本文
    participants = data.get("participants", [])

    # 出席者は1行に統合（会社別を1セルにまとめる）
    ROW_DATE     = 0
    ROW_PLACE    = 1
    ROW_PART     = 2   # 出席者（1行）
    ROW_MATERIAL = 3
    ROW_AGENDA   = 4
    ROW_HDR      = 5
    ROW_BODY     = 6
    TOTAL_ROWS   = 7

    info_tbl = doc.add_table(rows=TOTAL_ROWS, cols=3)
    info_tbl.style = "Table Grid"
    _set_col_width(info_tbl, [COL0_WIDTH, COL1_WIDTH, COL2_WIDTH])

    # --- 開催日時（列1+2をマージ）---
    _merge_h(info_tbl, ROW_DATE, 1, 2)
    _cell_para(info_tbl.cell(ROW_DATE, 0), "開催日時", bold=True)
    _cell_para(info_tbl.cell(ROW_DATE, 1), data.get("date", "") or "")

    # --- 場所（列1+2をマージ）---
    _merge_h(info_tbl, ROW_PLACE, 1, 2)
    _cell_para(info_tbl.cell(ROW_PLACE, 0), "場所", bold=True)
    _cell_para(info_tbl.cell(ROW_PLACE, 1), data.get("location", "") or "")

    # --- 出席者（1行に統合、列1+2をマージして全員を記載）---
    _merge_h(info_tbl, ROW_PART, 1, 2)
    _cell_para(info_tbl.cell(ROW_PART, 0), "出席者\n（敬称略）", bold=True)
    part_cell = info_tbl.cell(ROW_PART, 1)
    part_cell.paragraphs[0].clear()
    if participants:
        lines = []
        for p in participants:
            org = p.get("organization", "")
            names = p.get("names", "")
            if org and names:
                lines.append(f"{org}　{names}")
            elif org:
                lines.append(org)
            elif names:
                lines.append(names)
        part_run = part_cell.paragraphs[0].add_run("\n".join(lines))
        _set_font(part_run, 10)
    else:
        part_run = part_cell.paragraphs[0].add_run("")
        _set_font(part_run, 10)

    # --- 資料（列1+2をマージ）---
    _merge_h(info_tbl, ROW_MATERIAL, 1, 2)
    _cell_para(info_tbl.cell(ROW_MATERIAL, 0), "資料", bold=True)

    materials = data.get("materials", [])
    mat_cell = info_tbl.cell(ROW_MATERIAL, 1)
    _cell_para(mat_cell, "")
    mat_cell.paragraphs[0].clear()
    if materials:
        mat_para = mat_cell.paragraphs[0]
        mat_run = mat_para.add_run("\n".join(f"・{m}" for m in materials))
        _set_font(mat_run, 10)
    else:
        mat_run = mat_cell.paragraphs[0].add_run("")
        _set_font(mat_run, 10)

    # --- 議事次第（列1+2をマージ）---
    _merge_h(info_tbl, ROW_AGENDA, 1, 2)
    _cell_para(info_tbl.cell(ROW_AGENDA, 0), "議事次第", bold=True)
    agenda_summary = data.get("agenda_summary", [])
    ag_cell = info_tbl.cell(ROW_AGENDA, 1)
    ag_cell.paragraphs[0].clear()
    if agenda_summary:
        ag_run = ag_cell.paragraphs[0].add_run("\n".join(agenda_summary))
        _set_font(ag_run, 10)
    else:
        ag_run = ag_cell.paragraphs[0].add_run("")
        _set_font(ag_run, 10)

    # --- 議事内容ヘッダー（3列マージ）---
    _merge_full(info_tbl, ROW_HDR)
    _cell_para(info_tbl.cell(ROW_HDR, 0), "議事内容", bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- 本文エリア（3列マージ）---
    _merge_full(info_tbl, ROW_BODY)
    body_cell = info_tbl.cell(ROW_BODY, 0)
    # 既存の空段落をクリア
    body_cell.paragraphs[0].clear()
    body_para0 = body_cell.paragraphs[0]
    body_para0.paragraph_format.space_after = Pt(0)

    def _body_heading(text: str):
        """本文エリアの大見出しを追加。"""
        p = body_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        _set_font(r, 10, bold=True)

    def _body_bullet(text: str, indent: float = 0.5):
        """本文エリアの箇条書きを追加。"""
        p = body_cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"・{text}")
        _set_font(r, 10)

    def _body_text(text: str, indent: float = 0.0, bold: bool = False):
        """本文エリアの通常テキストを追加。"""
        p = body_cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        _set_font(r, 10, bold=bold)

    # ── １．議題ごとの議論 ──
    agenda_items = data.get("agenda_items", [])
    if agenda_items:
        _body_heading("１．議題")
        for item in agenda_items:
            _body_text(f"▶ {item.get('title', '')}", indent=0.3, bold=True)
            _body_text(item.get("discussion", ""), indent=0.8)

    # ── ２．決定事項・宿題事項 ──
    decisions = data.get("decisions", [])
    _body_heading("２．決定事項・宿題事項,etc.")
    if decisions:
        for d in decisions:
            _body_bullet(d)
    else:
        _body_bullet("（なし）")

    # ── ３．次回設定 ──
    _body_heading("３．次回設定、その他,etc.")
    next_date = data.get("next_date", "") or ""
    next_location = data.get("next_location", "") or ""
    if next_date:
        _body_text(f"日時：{next_date}", indent=0.3)
    if next_location:
        _body_text(f"場所：{next_location}", indent=0.3)
    next_topics = data.get("next_topics", [])
    if next_topics:
        for t in next_topics:
            _body_bullet(t)

    # 「以上」
    p_end = body_cell.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_end.paragraph_format.space_before = Pt(8)
    r_end = p_end.add_run("以上")
    _set_font(r_end, 10)

    # バイト列として返す
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if __name__ == "__main__":
    # 動作確認用サンプル実行
    sample_data = {
        "meeting_name": "第3回プロジェクト定例打合せ",
        "date": "令和8年7月25日（金）14:00～15:00",
        "location": "本社会議室A",
        "participants": [
            {"organization": "○○株式会社（事務局）", "names": "田中、山田"},
            {"organization": "△△株式会社（事業協力者）", "names": "佐藤"},
        ],
        "materials": [
            "第3回資料1 システム設計案（田中作成）",
            "第3回資料2 開発スケジュール（山田作成）",
        ],
        "agenda_summary": [
            "１．前回アクションの確認",
            "２．システム設計について",
            "３．開発環境の決定",
            "４．次回スケジュール確認",
        ],
        "agenda_items": [
            {
                "title": "前回アクションの確認",
                "discussion": "山田より仕様書が80%完成していることを報告。今週中に完成予定。"
            },
            {
                "title": "システム設計について",
                "discussion": "データベースはPostgreSQLを採用することで合意。パフォーマンス要件を満たすと判断。"
            },
            {
                "title": "開発環境の決定",
                "discussion": "開発環境はDockerで統一することに決定。各メンバーの環境差異を解消する。"
            }
        ],
        "decisions": [
            "データベースはPostgreSQLを採用する",
            "開発環境はDockerで統一する",
        ],
        "next_topics": [
            "詳細設計のレビュー",
            "テスト方針の策定",
        ],
        "next_date": "令和8年7月29日（水）14:00～",
        "next_location": "本社会議室A",
    }

    output_bytes = build_minutes_docx(sample_data, district="○○地区")
    output_path = "sample_minutes.docx"
    with open(output_path, "wb") as f:
        f.write(output_bytes)
    print(f"Word ファイルを生成しました: {output_path} ({len(output_bytes):,} bytes)")
